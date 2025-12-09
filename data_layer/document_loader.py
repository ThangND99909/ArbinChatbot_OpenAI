import fitz  # PyMuPDF
import PyPDF2
from docx import Document
import os
import json
import hashlib
import time
import signal
from typing import List, Dict, Tuple, Any, Optional, Callable
import logging
from datetime import datetime
import re
from pathlib import Path
import traceback
from concurrent.futures import ThreadPoolExecutor, as_completed

# Thêm import cho OCR
try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    pytesseract = None
    Image = None

logger = logging.getLogger(__name__)


class EnhancedDocumentProcessor:
    """
    Trình xử lý tài liệu nâng cao với:
    - PDF / DOCX / TXT / CSV / JSON / MD
    - Hybrid OCR (song song, adaptive)
    - Cache file đã xử lý
    - Quản lý bộ nhớ và timeout
    """

    def __init__(
        self,
        documents_dir: str = "./documents",
        data_manager=None,
        enable_ocr: bool = True,
        ocr_lang: str = "vie+eng",
        ocr_threshold: int = 50,
        ocr_max_pages: int = 50,
        ocr_timeout: int = 30,
        batch_size: int = 20,
        max_retries: int = 2,
        log_level=logging.INFO,
        log_file: Optional[str] = None
    ):
        self.documents_dir = Path(documents_dir)
        self.data_manager = data_manager
        self.documents_dir.mkdir(parents=True, exist_ok=True)

        # OCR Configuration
        self.enable_ocr = enable_ocr and OCR_AVAILABLE
        self.ocr_lang = ocr_lang
        self.ocr_threshold = ocr_threshold  # % trang scanned threshold
        self.ocr_max_pages = ocr_max_pages  # max pages for OCR in large files
        self.ocr_timeout = ocr_timeout  # timeout per page in seconds
        self.batch_size = batch_size  # batch size for memory management
        self.max_retries = max_retries  # retry attempts for failed processing

        if enable_ocr and not OCR_AVAILABLE:
            logger.warning(
                "OCR requested but pytesseract or PIL not installed. Install with: pip install pytesseract pillow"
            )

        # Cache processed files
        self.processed_files_info: Dict[str, Dict] = {}
        self.processed_log = self.documents_dir / "processed_files.json"
        self.load_processed_files()

        # File handlers
        self.file_handlers = {
            '.pdf': self.process_pdf,
            '.docx': self.process_docx,
            '.txt': self.process_text_file,
            '.md': self.process_text_file,
            '.rtf': self.process_text_file,
            '.csv': self.process_text_file,
            '.json': self.process_text_file
        }

        # Setup logging
        self.setup_logging(level=log_level, log_file=log_file)

    # =====================================================
    # LOGGING SETUP
    # =====================================================
    def setup_logging(self, level=logging.INFO, log_file=None):
        """Cấu hình logging chi tiết"""
        logger = logging.getLogger(__name__)
        logger.setLevel(level)

        # Clear existing handlers
        logger.handlers.clear()

        # Console handler
        console_handler = logging.StreamHandler()
        console_handler.setLevel(level)
        console_format = logging.Formatter(
            '%(asctime)s - %(name)s - %(levelname)s - %(message)s'
        )
        console_handler.setFormatter(console_format)
        logger.addHandler(console_handler)

        # File handler (nếu có)
        if log_file:
            file_handler = logging.FileHandler(log_file, encoding='utf-8')
            file_handler.setLevel(level)
            file_format = logging.Formatter(
                '%(asctime)s - %(name)s - %(levelname)s - %(filename)s:%(lineno)d - %(message)s'
            )
            file_handler.setFormatter(file_format)
            logger.addHandler(file_handler)

        return logger

    # =====================================================
    # CONFIGURATION VALIDATION
    # =====================================================
    def validate_configuration(self) -> Dict[str, bool]:
        """Validate hệ thống có đầy đủ dependencies không"""
        validation = {
            "pymupdf": fitz is not None,
            "pypdf2": PyPDF2 is not None,
            "python_docx": Document is not None,
            "pytesseract": OCR_AVAILABLE,
            "pillow": OCR_AVAILABLE,
            "documents_dir_exists": self.documents_dir.exists(),
            "documents_dir_writable": os.access(self.documents_dir, os.W_OK),
        }

        # Check tesseract command
        if OCR_AVAILABLE:
            try:
                pytesseract.get_tesseract_version()
                validation["tesseract_installed"] = True
                validation["tesseract_version"] = str(pytesseract.get_tesseract_version())
            except Exception as e:
                validation["tesseract_installed"] = False
                validation["tesseract_error"] = str(e)

        # Check OCR language
        if OCR_AVAILABLE and self.enable_ocr:
            try:
                # Test OCR with a simple command
                pytesseract.get_languages(config='')
                validation["ocr_languages_available"] = True
            except Exception as e:
                validation["ocr_languages_available"] = False
                validation["ocr_languages_error"] = str(e)

        return validation

    # =====================================================
    # CORE FILE TRACKING
    # =====================================================
    def load_processed_files(self):
        """Load cache of processed files"""
        if self.processed_log.exists():
            try:
                with open(self.processed_log, 'r', encoding='utf-8') as f:
                    self.processed_files_info = json.load(f)
                logger.info(f"Loaded {len(self.processed_files_info)} processed files info")
            except Exception as e:
                logger.warning(f"Could not load processed_files.json: {e}")
                self.processed_files_info = {}

    def save_processed_files(self):
        """Save cache of processed files"""
        try:
            with open(self.processed_log, 'w', encoding='utf-8') as f:
                json.dump(self.processed_files_info, f, indent=2, ensure_ascii=False)
            logger.debug(f"Saved processed file index: {len(self.processed_files_info)} entries")
        except Exception as e:
            logger.error(f"Error saving processed files: {e}")

    # =====================================================
    # FILE HASH & CHECKING
    # =====================================================
    def get_content_hash(self, content: str) -> str:
        """Generate hash from content"""
        return hashlib.md5(content.encode('utf-8')).hexdigest()

    def get_file_hash(self, file_path: str) -> str:
        """Generate hash from file metadata"""
        try:
            stat = os.stat(file_path)
            key = f"{file_path}_{stat.st_size}_{stat.st_mtime}"
            return hashlib.md5(key.encode('utf-8')).hexdigest()
        except Exception:
            return hashlib.md5(file_path.encode('utf-8')).hexdigest()

    def should_process_file(self, file_path: str) -> Tuple[bool, str, Dict]:
        """Check if file needs processing"""
        if not os.path.exists(file_path):
            return False, "File not found", {}
        
        key = str(Path(file_path).resolve())
        if key not in self.processed_files_info:
            return True, "New file", {}

        info = self.processed_files_info[key]
        try:
            current_mtime = os.path.getmtime(file_path)
            current_hash = self.get_file_hash(file_path)
            
            if current_mtime > info.get("file_mtime", 0):
                return True, "Modified file", info
            elif current_hash != info.get("file_hash", ""):
                return True, "Hash changed", info
            else:
                return False, "Unchanged", info
        except Exception:
            return True, "Error checking modification", info

    # =====================================================
    # OCR UTILITIES - Parallel with Timeout
    # =====================================================
    def extract_page_ocr_with_timeout(self, page_num: int, doc: fitz.Document) -> Tuple[int, str]:
        """Trích xuất OCR cho 1 trang PDF với timeout"""
        try:
            # Setup timeout (Unix only)
            def timeout_handler(signum, frame):
                raise TimeoutError(f"OCR timeout after {self.ocr_timeout} seconds")
            
            # Set timeout handler
            original_handler = signal.signal(signal.SIGALRM, timeout_handler)
            signal.alarm(self.ocr_timeout)
            
            try:
                page = doc.load_page(page_num)
                mat = fitz.Matrix(300/72, 300/72)
                pix = page.get_pixmap(matrix=mat)
                
                from io import BytesIO
                img = Image.open(BytesIO(pix.tobytes("ppm")))
                text = pytesseract.image_to_string(img, lang=self.ocr_lang, config='--psm 3 --oem 3')
                img.close()
                pix = None  # Giải phóng bộ nhớ
                
                return page_num, text.strip()
            finally:
                # Reset alarm
                signal.alarm(0)
                signal.signal(signal.SIGALRM, original_handler)
                
        except TimeoutError as e:
            logger.warning(f"OCR timeout on page {page_num}: {e}")
            return page_num, f"[OCR_TIMEOUT_AFTER_{self.ocr_timeout}s]"
        except Exception as e:
            logger.error(f"OCR error on page {page_num}: {e}")
            return page_num, ""

    def extract_text_with_ocr_parallel(self, pdf_path: str, max_pages: int = None) -> Tuple[str, int]:
        """OCR song song, giới hạn trang lớn"""
        if not self.enable_ocr:
            return "", 0

        max_pages = max_pages or self.ocr_max_pages
        doc = fitz.open(pdf_path)
        num_pages = len(doc)
        pages_to_process = list(range(min(num_pages, max_pages)))
        
        # Limit workers based on available resources
        num_workers = min(len(pages_to_process), os.cpu_count() or 2, 4)  # Max 4 workers
        logger.info(f"OCR parallel processing {len(pages_to_process)} pages with {num_workers} workers")

        results = [None] * len(pages_to_process)
        with ThreadPoolExecutor(max_workers=num_workers) as executor:
            future_to_page = {executor.submit(self.extract_page_ocr_with_timeout, p, doc): p 
                            for p in pages_to_process}
            
            for future in as_completed(future_to_page):
                try:
                    page_num, text = future.result(timeout=self.ocr_timeout + 5)
                    results[page_num] = f"--- Trang {page_num + 1} ---\n{text}" if text else ""
                except Exception as e:
                    page_num = future_to_page[future]
                    logger.error(f"Failed to process page {page_num}: {e}")
                    results[page_num] = ""

        doc.close()
        combined_text = "\n\n".join([r for r in results if r])
        return combined_text, len([r for r in results if r])

    def _process_ocr_batch(self, pdf_path: str, page_numbers: List[int]) -> str:
        """Process a batch of pages with OCR"""
        if not page_numbers:
            return ""
        
        try:
            doc = fitz.open(pdf_path)
            results = []
            
            for page_num in page_numbers:
                try:
                    page = doc.load_page(page_num)
                    mat = fitz.Matrix(300/72, 300/72)
                    pix = page.get_pixmap(matrix=mat)
                    
                    from io import BytesIO
                    img = Image.open(BytesIO(pix.tobytes("ppm")))
                    text = pytesseract.image_to_string(img, lang=self.ocr_lang, config='--psm 3 --oem 3')
                    img.close()
                    
                    if text.strip():
                        results.append(f"--- Trang {page_num + 1} ---\n{text.strip()}")
                except Exception as e:
                    logger.warning(f"OCR failed for page {page_num}: {e}")
                    continue
            
            doc.close()
            return "\n\n".join(results)
        except Exception as e:
            logger.error(f"Batch OCR processing failed: {e}")
            return ""

    # =====================================================
    # PDF PROCESSING
    # =====================================================
    def is_pdf_scanned(self, pdf_path: str, threshold: float = 0.1) -> Tuple[bool, float]:
        """Check if PDF is scanned/image-based"""
        try:
            with fitz.open(pdf_path) as doc:
                total_pages = len(doc)
                if total_pages == 0:
                    return False, 0.0
                
                scanned_pages = 0
                sample_pages = min(total_pages, 10)
                
                for page_num in range(sample_pages):
                    try:
                        page = doc.load_page(page_num)
                        text = page.get_text("text").strip()
                        
                        # Check for very little text
                        if len(text) < 50:
                            scanned_pages += 1
                    except Exception:
                        scanned_pages += 1  # Assume scanned if error
                
                scan_ratio = scanned_pages / sample_pages
                return scan_ratio > threshold, scan_ratio
        except Exception as e:
            logger.error(f"Error checking if PDF is scanned: {e}")
            return False, 0.0

    def extract_pdf_hybrid_parallel(self, pdf_path: str) -> Tuple[str, int, Dict]:
        """Hybrid OCR với parallel processing và quản lý bộ nhớ"""
        try:
            doc = fitz.open(pdf_path)
            num_pages = len(doc)
            
            # Giới hạn số trang xử lý cùng lúc để tránh quá tải bộ nhớ
            batch_size = min(self.batch_size, num_pages)
            all_text_parts = []
            ocr_pages = []
            normal_pages = []
            
            # Phân tích từng batch
            for start in range(0, num_pages, batch_size):
                end = min(start + batch_size, num_pages)
                
                # Xử lý batch hiện tại
                for page_num in range(start, end):
                    try:
                        page = doc.load_page(page_num)
                        normal_text = page.get_text("text").strip()
                        
                        if len(normal_text) < 50:  # Threshold for considering as scanned
                            ocr_pages.append(page_num)
                        else:
                            normal_pages.append(page_num)
                            all_text_parts.append(f"--- Trang {page_num + 1} ---\n{normal_text}")
                    except Exception as e:
                        logger.warning(f"Error processing page {page_num}: {e}")
                        ocr_pages.append(page_num)  # Assume needs OCR if error
                
                # Giải phóng bộ nhớ sau mỗi batch
                doc._cleanup()
            
            # Xử lý OCR theo batch để tránh quá tải
            ocr_text_parts = []
            if ocr_pages and self.enable_ocr:
                logger.info(f"Processing {len(ocr_pages)} pages with OCR")
                
                # Chia nhỏ OCR pages thành các batch
                ocr_batch_size = min(10, self.ocr_max_pages // 2)
                for i in range(0, len(ocr_pages), ocr_batch_size):
                    batch = ocr_pages[i:i + ocr_batch_size]
                    batch_text = self._process_ocr_batch(pdf_path, batch)
                    if batch_text:
                        ocr_text_parts.append(batch_text)
                    
                    # Small delay between batches to prevent resource exhaustion
                    if i + ocr_batch_size < len(ocr_pages):
                        time.sleep(0.1)
            
            all_text_parts.extend(ocr_text_parts)
            
            method_info = {
                "extraction_method": "hybrid_parallel_ocr",
                "normal_pages": normal_pages,
                "ocr_pages": ocr_pages,
                "ocr_ratio": len(ocr_pages)/num_pages if num_pages > 0 else 0,
                "batches_processed": (num_pages + batch_size - 1) // batch_size
            }
            
            doc.close()
            return "\n\n".join(all_text_parts), num_pages, method_info
            
        except Exception as e:
            logger.error(f"Hybrid parallel OCR error: {e}")
            return "", 0, {}

    def extract_pdf_text_fitz(self, file_path: str) -> Tuple[str, int]:
        """Extract text using PyMuPDF"""
        text_blocks = []
        with fitz.open(file_path) as doc:
            num_pages = len(doc)
            for page_num, page in enumerate(doc):
                try:
                    page_text = page.get_text("text")
                    if page_text.strip():
                        text_blocks.append(f"--- Trang {page_num + 1} ---\n{page_text.strip()}")
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num}: {e}")
                    text_blocks.append(f"--- Trang {page_num + 1} ---\n[ERROR_EXTRACTING_TEXT]")
        
        return "\n\n".join(text_blocks), num_pages

    def extract_pdf_metadata(self, reader: PyPDF2.PdfReader) -> Dict[str, str]:
        """Extract PDF metadata"""
        meta = {}
        try:
            pdf_meta = reader.metadata
            if pdf_meta:
                for k, v in pdf_meta.items():
                    key = k.replace("/", "").lower()
                    meta[key] = str(v)
        except Exception as e:
            logger.debug(f"Error extracting PDF metadata: {e}")
        
        return meta

    def process_pdf(self, file_path: str, force_reprocess: bool = False) -> Dict[str, Any]:
        """Process PDF file with hybrid approach"""
        text = ""
        metadata = {}
        num_pages = 0
        extraction_method = "fitz"
        
        try:
            # Check if PDF is scanned
            is_scanned, scan_ratio = self.is_pdf_scanned(file_path)
            logger.info(f"PDF scan check: is_scanned={is_scanned}, ratio={scan_ratio:.2f}")
            
            # Try standard extraction first
            text, num_pages = self.extract_pdf_text_fitz(file_path)
            text_length = len(text.strip())
            
            # Decide extraction method
            if text_length < 100 or (is_scanned and scan_ratio > self.ocr_threshold/100):
                if self.enable_ocr:
                    logger.info(f"Using hybrid parallel OCR for {file_path}")
                    text, num_pages, ocr_info = self.extract_pdf_hybrid_parallel(file_path)
                    extraction_method = "hybrid_parallel_ocr"
                    metadata.update(ocr_info)
                else:
                    logger.info(f"Low text but OCR disabled, using PyPDF2 fallback")
                    extraction_method = "PyPDF2_fallback"
                    raise Exception("Low text, OCR disabled")
            else:
                extraction_method = "fitz_standard"
                logger.info(f"Using standard Fitz extraction for {file_path}")
        
        except Exception as e:
            logger.warning(f"Fitz extraction failed, trying PyPDF2 fallback: {e}")
            # Fallback to PyPDF2
            try:
                with open(file_path, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    metadata = self.extract_pdf_metadata(reader)
                    num_pages = len(reader.pages)
                    
                    text_parts = []
                    for page_num, page in enumerate(reader.pages):
                        try:
                            page_text = page.extract_text() or ""
                            page_text = re.sub(r"\s+", " ", page_text).strip()
                            if page_text:
                                text_parts.append(f"--- Trang {page_num + 1} ---\n{page_text}")
                        except Exception as page_error:
                            logger.warning(f"Error extracting page {page_num}: {page_error}")
                    
                    text = "\n\n".join(text_parts)
                    extraction_method = "PyPDF2_fallback"
                    
            except Exception as e2:
                logger.error(f"Error reading PDF {file_path} with PyPDF2 fallback: {e2}")
                return {}

        # Ensure metadata
        if not metadata.get("title"):
            # Try to extract title from text
            lines = text.splitlines()
            for line in lines[:10]:  # Check first 10 lines
                if 10 < len(line.strip()) < 150:
                    metadata["title"] = line.strip()[:100]
                    break
            if not metadata.get("title"):
                metadata["title"] = os.path.basename(file_path)

        metadata["num_pages"] = num_pages
        metadata["extraction_method"] = extraction_method
        metadata["scan_ratio"] = scan_ratio if 'scan_ratio' in locals() else 0.0
        
        content_hash = self.get_content_hash(text)
        file_hash = self.get_file_hash(file_path)
        
        return self._make_result(file_path, text, "pdf", metadata, content_hash, file_hash)

    # =====================================================
    # DOCX PROCESSING
    # =====================================================
    def extract_docx_metadata(self, doc: Document) -> Dict[str, str]:
        """Extract DOCX metadata"""
        meta = {}
        try:
            core = doc.core_properties
            meta = {
                "title": core.title or "",
                "author": core.author or "",
                "subject": core.subject or "",
                "created": str(core.created or ""),
                "modified": str(core.modified or ""),
                "last_modified_by": core.last_modified_by or "",
                "keywords": core.keywords or "",
                "category": core.category or "",
                "comments": core.comments or "",
                "identifier": core.identifier or "",
                "language": core.language or "",
                "version": core.version or "",
                "content_status": core.content_status or "",
                "revision": str(core.revision or "")
            }
        except Exception as e:
            logger.debug(f"Error extracting DOCX metadata: {e}")
        
        return meta

    def process_docx(self, file_path: str, force_reprocess: bool = False) -> Dict[str, Any]:
        """Process DOCX file"""
        try:
            doc = Document(file_path)
            
            # Extract paragraphs
            paragraphs = []
            for p in doc.paragraphs:
                if p.text.strip():
                    paragraphs.append(p.text.strip())
            
            # Extract tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        tables_text.append(" | ".join(row_text))
            
            # Combine content
            text_parts = []
            if paragraphs:
                text_parts.append("\n\n".join(paragraphs))
            if tables_text:
                text_parts.append("--- TABLES ---\n" + "\n".join(tables_text))
            
            text = "\n\n".join(text_parts)
            metadata = self.extract_docx_metadata(doc)
            metadata["num_paragraphs"] = len(paragraphs)
            metadata["num_tables"] = len(doc.tables)
            
            if not metadata.get("title"):
                metadata["title"] = os.path.basename(file_path)
                
        except Exception as e:
            logger.error(f"Error reading DOCX {file_path}: {e}")
            return {}

        content_hash = self.get_content_hash(text)
        file_hash = self.get_file_hash(file_path)
        
        return self._make_result(file_path, text, "docx", metadata, content_hash, file_hash)

    # =====================================================
    # TEXT FILES PROCESSING
    # =====================================================
    def process_text_file(self, file_path: str, force_reprocess: bool = False) -> Dict[str, Any]:
        """Process text-based files (txt, md, rtf, csv, json)"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
            text = None
            
            for encoding in encodings:
                try:
                    with open(file_path, "r", encoding=encoding) as f:
                        text = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if text is None:
                # Last resort: read as binary and decode with errors ignored
                with open(file_path, "rb") as f:
                    text = f.read().decode('utf-8', errors='ignore')
            
            metadata = {
                "num_lines": len(text.splitlines()),
                "title": os.path.basename(file_path),
                "encoding": encoding if 'encoding' in locals() else 'unknown',
                "file_type": Path(file_path).suffix.lower()
            }
            
        except Exception as e:
            logger.error(f"Error reading text file {file_path}: {e}")
            return {}

        content_hash = self.get_content_hash(text)
        file_hash = self.get_file_hash(file_path)
        
        return self._make_result(file_path, text, "text", metadata, content_hash, file_hash)

    # =====================================================
    # CORE RESULT BUILDER
    # =====================================================
    def _make_result(self, file_path, text, source, metadata, content_hash, file_hash):
        """Create standardized result dictionary"""
        try:
            file_size = os.path.getsize(file_path)
            file_mtime = os.path.getmtime(file_path)
        except Exception:
            file_size = 0
            file_mtime = 0
        
        # Calculate word count (simple)
        words = re.findall(r'\b\w+\b', text)
        word_count = len(words)
        
        meta = {
            **metadata,
            "file_size": file_size,
            "file_mtime": file_mtime,
            "word_count": word_count,
            "content_length": len(text),
            "character_count": len(text),
            "line_count": len(text.splitlines()),
        }
        
        return {
            "file_path": str(file_path),
            "file_name": os.path.basename(file_path),
            "content": text,
            "source": source,
            "source_type": "document",
            "content_hash": content_hash,
            "file_hash": file_hash,
            "metadata": meta,
            "processed_at": datetime.now().isoformat(),
            "processing_version": "2.0"
        }

    # =====================================================
    # MAIN PROCESSING LOOP WITH RETRY
    # =====================================================
    def process_file(self, file_path: str, force_reprocess: bool = False) -> Dict[str, Any]:
        """Process file with retry mechanism"""
        ext = Path(file_path).suffix.lower()
        if ext not in self.file_handlers:
            logger.debug(f"Skipping unsupported file: {file_path}")
            return {}

        should, reason, prev = self.should_process_file(file_path)
        if not should and not force_reprocess:
            logger.info(f"Skipping {os.path.basename(file_path)} - {reason}")
            return {"cached": True, **prev}

        logger.info(f"Processing {os.path.basename(file_path)} - {reason}")
        handler = self.file_handlers[ext]
        
        # Retry logic
        for attempt in range(self.max_retries + 1):
            try:
                result = handler(file_path, force_reprocess)
                
                if result and result.get("content"):
                    # Validate result
                    if len(result.get("content", "").strip()) < 10:
                        logger.warning(f"Very short content extracted ({len(result['content'])} chars)")
                    
                    key = str(Path(file_path).resolve())
                    self.processed_files_info[key] = {
                        "file_mtime": os.path.getmtime(file_path),
                        "file_hash": result["file_hash"],
                        "content_hash": result["content_hash"],
                        "source": result["source"],
                        "metadata": result["metadata"],
                        "processed_at": result["processed_at"],
                        "content_length": len(result.get("content", "")),
                    }
                    return result
                elif attempt < self.max_retries:
                    logger.warning(f"Empty content, retry {attempt + 1}/{self.max_retries}")
                    time.sleep(1)  # Wait before retry
                    
            except Exception as e:
                logger.error(f"Attempt {attempt + 1} failed: {e}")
                if attempt == self.max_retries:
                    logger.error(f"Failed to process {file_path} after {self.max_retries} attempts")
                    # Save error info
                    key = str(Path(file_path).resolve())
                    self.processed_files_info[key] = {
                        "file_mtime": os.path.getmtime(file_path),
                        "error": str(e),
                        "processed_at": datetime.now().isoformat(),
                        "failed": True
                    }
                    return {}
                time.sleep(1)  # Wait before retry
        
        return {}

    def process_all_documents(self, force_reprocess: bool = False, 
                            callback: Optional[Callable] = None) -> List[Dict[str, Any]]:
        """Xử lý tất cả document với progress callback"""
        files = [
            p for p in self.documents_dir.rglob("*")
            if p.suffix.lower() in self.file_handlers and not p.name.startswith("~$")
        ]
        
        if not files:
            logger.warning(f"No documents found in {self.documents_dir}")
            return []
        
        logger.info(f"Found {len(files)} supported files in {self.documents_dir}")
        results = []
        failed_files = []
        
        for i, file in enumerate(files):
            file_path = str(file)
            
            # Progress callback
            if callback:
                callback(i, len(files), file_path)
            
            logger.debug(f"Processing file {i+1}/{len(files)}: {file.name}")
            
            try:
                doc = self.process_file(file_path, force_reprocess)
                if doc:
                    if not doc.get("cached", False):
                        content_len = len(doc.get("content", ""))
                        if content_len < 50:
                            logger.warning(f"{file.name} extracted very little text ({content_len} chars)")
                        else:
                            logger.info(f"Successfully processed {file.name} ({content_len} chars)")
                    
                    results.append(doc)
                else:
                    failed_files.append(file_path)
                    logger.error(f"Failed to process {file.name}")
                    
            except Exception as e:
                failed_files.append(file_path)
                logger.error(f"Unexpected error processing {file.name}: {e}")
                logger.debug(traceback.format_exc())
        
        # Save cache
        self.save_processed_files()
        
        # Save to data manager if available
        if self.data_manager:
            try:
                self.data_manager.save_raw_data(results, "enhanced_document_processor")
                
                summary = {
                    "total_files": len(files),
                    "successful": len(results),
                    "failed": len(failed_files),
                    "failed_files": failed_files,
                    "dir": str(self.documents_dir),
                    "timestamp": datetime.now().isoformat(),
                    "stats": self.get_detailed_stats()
                }
                
                self.data_manager.export_for_inspection(summary, "doc_processing_summary")
                logger.info(f"Saved processing summary to data manager")
                
            except Exception as e:
                logger.error(f"Error saving to data manager: {e}")
        
        logger.info(f"Processing complete: {len(results)} successful, {len(failed_files)} failed")
        return results

    # =====================================================
    # STATISTICS AND MONITORING
    # =====================================================
    def get_detailed_stats(self) -> Dict[str, Any]:
        """Lấy thống kê chi tiết"""
        stats = {
            "total_files": len(self.processed_files_info),
            "by_type": {},
            "by_size": {"small": 0, "medium": 0, "large": 0},
            "ocr_stats": {
                "ocr_used": 0,
                "ocr_pages": 0,
                "hybrid_used": 0,
                "total_ocr_pages": 0
            },
            "content_stats": {
                "total_words": 0,
                "total_chars": 0,
                "avg_words_per_file": 0,
                "avg_chars_per_file": 0
            },
            "file_stats": {
                "cached": 0,
                "modified": 0,
                "new": 0
            }
        }
        
        for key, info in self.processed_files_info.items():
            # Skip failed files
            if info.get("failed"):
                continue
            
            # Thống kê theo type
            src = info.get("source", "unknown")
            stats["by_type"][src] = stats["by_type"].get(src, 0) + 1
            
            # Thống kê theo size
            size = info.get("metadata", {}).get("file_size", 0)
            if size < 1024 * 1024:  # < 1MB
                stats["by_size"]["small"] += 1
            elif size < 10 * 1024 * 1024:  # < 10MB
                stats["by_size"]["medium"] += 1
            else:
                stats["by_size"]["large"] += 1
            
            # Thống kê OCR
            metadata = info.get("metadata", {})
            if metadata.get("extraction_method") == "hybrid_parallel_ocr":
                stats["ocr_stats"]["hybrid_used"] += 1
                ocr_pages = len(metadata.get("ocr_pages", []))
                stats["ocr_stats"]["ocr_pages"] += ocr_pages
                stats["ocr_stats"]["total_ocr_pages"] += ocr_pages
            elif metadata.get("extraction_method") == "ocr_only":
                stats["ocr_stats"]["ocr_used"] += 1
            
            # Thống kê content
            stats["content_stats"]["total_words"] += metadata.get("word_count", 0)
            stats["content_stats"]["total_chars"] += metadata.get("content_length", 0)
        
        # Tính trung bình
        valid_files = stats["total_files"]
        if valid_files > 0:
            stats["content_stats"]["avg_words_per_file"] = (
                stats["content_stats"]["total_words"] / valid_files
            )
            stats["content_stats"]["avg_chars_per_file"] = (
                stats["content_stats"]["total_chars"] / valid_files
            )
        
        return stats

    def get_document_stats(self) -> Dict[str, Any]:
        """Get basic statistics (backward compatibility)"""
        return self.get_detailed_stats()

    # =====================================================
    # UTILITIES
    # =====================================================
    def get_supported_extensions(self) -> List[str]:
        """Get list of supported file extensions"""
        return list(self.file_handlers.keys())

    def clear_processed_cache(self):
        """Clear processing cache"""
        self.processed_files_info.clear()
        if self.processed_log.exists():
            try:
                self.processed_log.unlink()
                logger.info("Cleared document processor cache")
            except Exception as e:
                logger.error(f"Error clearing cache: {e}")

    def set_ocr_language(self, lang: str):
        """Set OCR language"""
        self.ocr_lang = lang
        logger.info(f"OCR language set to: {lang}")

    def enable_disable_ocr(self, enable: bool):
        """Enable or disable OCR"""
        if enable and not OCR_AVAILABLE:
            logger.warning("Cannot enable OCR: pytesseract or PIL not installed")
            return
        
        old_status = self.enable_ocr
        self.enable_ocr = enable
        logger.info(f"OCR changed from {old_status} to {self.enable_ocr}")

    def get_processing_summary(self) -> Dict[str, Any]:
        """Get comprehensive processing summary"""
        validation = self.validate_configuration()
        stats = self.get_detailed_stats()
        
        return {
            "configuration": {
                "documents_dir": str(self.documents_dir),
                "enable_ocr": self.enable_ocr,
                "ocr_lang": self.ocr_lang,
                "ocr_threshold": self.ocr_threshold,
                "ocr_max_pages": self.ocr_max_pages,
                "ocr_timeout": self.ocr_timeout,
                "batch_size": self.batch_size,
                "max_retries": self.max_retries,
            },
            "validation": validation,
            "statistics": stats,
            "cache_info": {
                "cached_files": len(self.processed_files_info),
                "cache_file": str(self.processed_log),
                "cache_exists": self.processed_log.exists()
            },
            "system_info": {
                "python_version": os.sys.version,
                "platform": os.sys.platform,
                "cpu_count": os.cpu_count(),
                "timestamp": datetime.now().isoformat()
            }
        }


# Backward compatibility
DocumentProcessor = EnhancedDocumentProcessor


# Example usage
if __name__ == "__main__":
    # Setup basic logging
    logging.basicConfig(level=logging.INFO)
    
    # Create processor instance
    processor = EnhancedDocumentProcessor(
        documents_dir="./test_documents",
        enable_ocr=True,
        ocr_lang="vie+eng",
        ocr_threshold=30,
        ocr_max_pages=100,
        ocr_timeout=30,
        batch_size=10,
        max_retries=2,
        log_level=logging.INFO,
        log_file="document_processor.log"
    )
    
    # Validate configuration
    validation = processor.validate_configuration()
    print("Configuration validation:")
    for key, value in validation.items():
        print(f"  {key}: {value}")
    
    # Process all documents with progress callback
    def progress_callback(current, total, filename):
        percent = (current / total) * 100
        print(f"\rProcessing: {current}/{total} ({percent:.1f}%) - {Path(filename).name}", end="")
    
    print("\nStarting document processing...")
    results = processor.process_all_documents(
        force_reprocess=False,
        callback=progress_callback
    )
    
    print(f"\n\nProcessing complete. Processed {len(results)} documents.")
    
    # Show statistics
    stats = processor.get_processing_summary()
    print("\nProcessing Summary:")
    print(json.dumps(stats, indent=2, ensure_ascii=False))